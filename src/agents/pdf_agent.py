"""
McLeuker Agentic AI Platform - PDF/Word Execution Agent

Generates PDF and Word documents from structured data.
"""

import os
from typing import Optional, List, Dict, Any
from datetime import datetime
from io import BytesIO

from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    PageBreak, ListFlowable, ListItem
)
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

from src.agents.base_agent import BaseExecutionAgent
from src.models.schemas import (
    StructuredOutput,
    StructuredDocument,
    GeneratedFile,
    OutputFormat
)


class PDFAgent(BaseExecutionAgent):
    """
    PDF/Word Execution Agent
    
    Generates PDF and Word documents from structured document data.
    """
    
    @property
    def supported_formats(self) -> List[OutputFormat]:
        return [OutputFormat.PDF, OutputFormat.WORD]
    
    async def execute(
        self,
        structured_output: StructuredOutput,
        filename_prefix: Optional[str] = None
    ) -> List[GeneratedFile]:
        """
        Generate PDF and Word files from structured output.
        
        Args:
            structured_output: The structured data containing documents
            filename_prefix: Optional prefix for filenames
            
        Returns:
            List[GeneratedFile]: List of generated files
        """
        generated_files = []
        
        if not structured_output.documents:
            return generated_files
        
        for doc in structured_output.documents:
            # Generate PDF
            pdf_file = await self._generate_pdf(doc, filename_prefix)
            if pdf_file:
                generated_files.append(pdf_file)
            
            # Generate Word document
            word_file = await self._generate_word(doc, filename_prefix)
            if word_file:
                generated_files.append(word_file)
        
        return generated_files
    
    async def _generate_pdf(
        self,
        document: StructuredDocument,
        filename_prefix: Optional[str] = None
    ) -> Optional[GeneratedFile]:
        """Generate a PDF file from a structured document."""
        try:
            filename = self._generate_filename(filename_prefix, "pdf")
            filepath = self._get_file_path(filename)
            
            # Create PDF document
            doc = SimpleDocTemplate(
                filepath,
                pagesize=A4,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=72
            )
            
            # Get styles
            styles = getSampleStyleSheet()
            
            # Custom styles
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Title'],
                fontSize=24,
                spaceAfter=30,
                textColor=colors.HexColor('#1F4E79')
            )
            
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading1'],
                fontSize=16,
                spaceBefore=20,
                spaceAfter=10,
                textColor=colors.HexColor('#1F4E79')
            )
            
            subheading_style = ParagraphStyle(
                'CustomSubheading',
                parent=styles['Heading2'],
                fontSize=14,
                spaceBefore=15,
                spaceAfter=8,
                textColor=colors.HexColor('#2E75B6')
            )
            
            body_style = ParagraphStyle(
                'CustomBody',
                parent=styles['Normal'],
                fontSize=11,
                leading=16,
                alignment=TA_JUSTIFY,
                spaceAfter=12
            )
            
            # Build content
            story = []
            
            # Title
            story.append(Paragraph(document.title, title_style))
            
            # Metadata
            if document.metadata:
                meta_parts = []
                if document.metadata.get('author'):
                    meta_parts.append(f"Author: {document.metadata['author']}")
                if document.metadata.get('date'):
                    meta_parts.append(f"Date: {document.metadata['date']}")
                else:
                    meta_parts.append(f"Date: {datetime.now().strftime('%B %d, %Y')}")
                
                meta_text = " | ".join(meta_parts)
                meta_style = ParagraphStyle(
                    'Metadata',
                    parent=styles['Normal'],
                    fontSize=10,
                    textColor=colors.grey,
                    spaceAfter=20
                )
                story.append(Paragraph(meta_text, meta_style))
            
            story.append(Spacer(1, 20))
            
            # Sections
            for section in document.sections:
                # Section heading
                heading = section.get('heading', 'Section')
                story.append(Paragraph(heading, heading_style))
                
                # Section content
                content = section.get('content', '')
                if content:
                    # Split content into paragraphs
                    paragraphs = content.split('\n\n')
                    for para in paragraphs:
                        if para.strip():
                            story.append(Paragraph(para.strip(), body_style))
                
                # Subsections
                subsections = section.get('subsections', [])
                for subsection in subsections:
                    sub_heading = subsection.get('heading', '')
                    if sub_heading:
                        story.append(Paragraph(sub_heading, subheading_style))
                    
                    sub_content = subsection.get('content', '')
                    if sub_content:
                        paragraphs = sub_content.split('\n\n')
                        for para in paragraphs:
                            if para.strip():
                                story.append(Paragraph(para.strip(), body_style))
                
                story.append(Spacer(1, 10))
            
            # Build PDF
            doc.build(story)
            
            return self._create_generated_file(filename, OutputFormat.PDF)
            
        except Exception as e:
            print(f"Error generating PDF file: {e}")
            return None
    
    async def _generate_word(
        self,
        document: StructuredDocument,
        filename_prefix: Optional[str] = None
    ) -> Optional[GeneratedFile]:
        """Generate a Word document from a structured document."""
        try:
            filename = self._generate_filename(filename_prefix, "docx")
            filepath = self._get_file_path(filename)
            
            # Create Word document
            doc = Document()
            
            # Set document properties
            core_properties = doc.core_properties
            core_properties.title = document.title
            if document.metadata:
                if document.metadata.get('author'):
                    core_properties.author = document.metadata['author']
            
            # Add title
            title = doc.add_heading(document.title, level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add metadata
            if document.metadata:
                meta_parts = []
                if document.metadata.get('author'):
                    meta_parts.append(f"Author: {document.metadata['author']}")
                meta_parts.append(f"Date: {datetime.now().strftime('%B %d, %Y')}")
                
                meta_para = doc.add_paragraph(" | ".join(meta_parts))
                meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in meta_para.runs:
                    run.font.size = Pt(10)
                    run.font.italic = True
            
            doc.add_paragraph()  # Spacer
            
            # Add sections
            for section in document.sections:
                # Section heading
                heading = section.get('heading', 'Section')
                doc.add_heading(heading, level=1)
                
                # Section content
                content = section.get('content', '')
                if content:
                    paragraphs = content.split('\n\n')
                    for para in paragraphs:
                        if para.strip():
                            p = doc.add_paragraph(para.strip())
                            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                
                # Subsections
                subsections = section.get('subsections', [])
                for subsection in subsections:
                    sub_heading = subsection.get('heading', '')
                    if sub_heading:
                        doc.add_heading(sub_heading, level=2)
                    
                    sub_content = subsection.get('content', '')
                    if sub_content:
                        paragraphs = sub_content.split('\n\n')
                        for para in paragraphs:
                            if para.strip():
                                p = doc.add_paragraph(para.strip())
                                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            # Save document
            doc.save(filepath)
            
            return self._create_generated_file(filename, OutputFormat.WORD)
            
        except Exception as e:
            print(f"Error generating Word file: {e}")
            return None
    
    async def generate_simple_report(
        self,
        title: str,
        content: str,
        author: Optional[str] = "McLeuker AI",
        filename_prefix: Optional[str] = None
    ) -> List[GeneratedFile]:
        """
        Generate a simple report with just title and content.
        
        Args:
            title: Report title
            content: Report content
            author: Author name
            filename_prefix: Optional prefix for filename
            
        Returns:
            List of generated files
        """
        document = StructuredDocument(
            title=title,
            sections=[
                {
                    "heading": "Report",
                    "content": content
                }
            ],
            metadata={
                "author": author,
                "date": datetime.now().strftime('%Y-%m-%d')
            }
        )
        
        structured_output = StructuredOutput(documents=[document])
        return await self.execute(structured_output, filename_prefix)
