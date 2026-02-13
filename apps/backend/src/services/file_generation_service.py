"""
File Generation Service — Kimi-2.5 Style
=========================================

Generates actual downloadable files instead of dumping code in chat.
Supports: Excel, Word, PDF, PowerPoint, Code projects (multi-file folders)
"""

import asyncio
import os
import uuid
import shutil
import zipfile
from pathlib import Path
from typing import Dict, List, Optional, Any
from dataclasses import dataclass
from datetime import datetime
import logging

logger = logging.getLogger(__name__)

# File generation workspace
WORKSPACE_ROOT = Path("/tmp/mcleuker_files")
WORKSPACE_ROOT.mkdir(parents=True, exist_ok=True)


@dataclass
class GeneratedFile:
    """Represents a generated file."""
    filename: str
    filepath: str
    file_type: str  # excel, word, pdf, ppt, code, text
    size_bytes: int
    download_url: str
    created_at: datetime


@dataclass
class FileGenerationResult:
    """Result of file generation."""
    success: bool
    files: List[GeneratedFile]
    folder_path: Optional[str] = None
    zip_path: Optional[str] = None
    error: Optional[str] = None


class FileGenerationService:
    """
    Service for generating downloadable files from user requests.
    
    Features:
    - Execute code to generate files (not just return code)
    - Organize multi-file outputs into folders
    - Auto-zip folders for easy download
    - S3 upload for public URLs
    """
    
    def __init__(self, s3_client=None):
        """
        Args:
            s3_client: Optional S3 client for uploading files
        """
        self.s3_client = s3_client
        self.workspace_root = WORKSPACE_ROOT
    
    async def generate_excel(
        self,
        user_id: str,
        session_id: str,
        data: List[List[Any]],
        headers: Optional[List[str]] = None,
        sheet_name: str = "Sheet1",
        filename: Optional[str] = None,
        formatting: Optional[Dict] = None,
    ) -> FileGenerationResult:
        """
        Generate an Excel file from data.
        
        Args:
            user_id: User ID
            session_id: Session ID
            data: 2D list of data
            headers: Column headers
            sheet_name: Sheet name
            filename: Output filename (auto-generated if None)
            formatting: Optional formatting config
        
        Returns:
            FileGenerationResult with file path and download URL
        """
        try:
            import pandas as pd
            from openpyxl import Workbook
            from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
            from openpyxl.utils.dataframe import dataframe_to_rows
            from openpyxl.utils import get_column_letter
            
            # Create session workspace
            session_dir = self.workspace_root / user_id / session_id
            session_dir.mkdir(parents=True, exist_ok=True)
            
            # Generate filename
            if not filename:
                filename = f"generated_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
            
            filepath = session_dir / filename
            
            # Create DataFrame
            df = pd.DataFrame(data, columns=headers)
            
            # Create Excel workbook
            wb = Workbook()
            ws = wb.active
            ws.title = sheet_name
            
            # Write data
            for r_idx, row in enumerate(dataframe_to_rows(df, index=False, header=True), 1):
                for c_idx, value in enumerate(row, 1):
                    cell = ws.cell(row=r_idx, column=c_idx, value=value)
                    
                    # Header formatting
                    if r_idx == 1:
                        cell.font = Font(bold=True, color="FFFFFF")
                        cell.fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
                        cell.alignment = Alignment(horizontal='center', vertical='center')
                    else:
                        cell.alignment = Alignment(vertical='center')
                    
                    # Borders
                    cell.border = Border(
                        left=Side(style='thin'),
                        right=Side(style='thin'),
                        top=Side(style='thin'),
                        bottom=Side(style='thin')
                    )
            
            # Auto-adjust column widths
            for column in ws.columns:
                max_length = 0
                column_letter = get_column_letter(column[0].column)
                for cell in column:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(cell.value)
                    except:
                        pass
                adjusted_width = min(max_length + 2, 50)
                ws.column_dimensions[column_letter].width = adjusted_width
            
            # Apply custom formatting if provided
            if formatting:
                # TODO: Apply custom formatting
                pass
            
            # Save
            wb.save(str(filepath))
            
            # Get file size
            size_bytes = filepath.stat().st_size
            
            # Upload to S3 if available
            download_url = await self._upload_to_s3(filepath, user_id, session_id)
            
            generated_file = GeneratedFile(
                filename=filename,
                filepath=str(filepath),
                file_type="excel",
                size_bytes=size_bytes,
                download_url=download_url,
                created_at=datetime.now(),
            )
            
            logger.info(f"Generated Excel file: {filename} ({size_bytes} bytes)")
            
            return FileGenerationResult(
                success=True,
                files=[generated_file],
            )
            
        except Exception as e:
            logger.error(f"Excel generation failed: {e}")
            return FileGenerationResult(
                success=False,
                files=[],
                error=str(e),
            )
    
    async def generate_document(
        self,
        user_id: str,
        session_id: str,
        content: str,
        format: str = "pdf",  # pdf, docx, md
        filename: Optional[str] = None,
        styling: Optional[Dict] = None,
    ) -> FileGenerationResult:
        """
        Generate a document (PDF, Word, Markdown).
        
        Args:
            user_id: User ID
            session_id: Session ID
            content: Document content (Markdown or plain text)
            format: Output format (pdf, docx, md)
            filename: Output filename
            styling: Optional styling config
        
        Returns:
            FileGenerationResult
        """
        try:
            # Create session workspace
            session_dir = self.workspace_root / user_id / session_id
            session_dir.mkdir(parents=True, exist_ok=True)
            
            # Generate filename
            if not filename:
                filename = f"document_{datetime.now().strftime('%Y%m%d_%H%M%S')}.{format}"
            
            filepath = session_dir / filename
            
            if format == "md":
                # Save as Markdown
                filepath.write_text(content, encoding='utf-8')
            
            elif format == "pdf":
                # Convert Markdown to PDF using WeasyPrint
                from weasyprint import HTML, CSS
                from markdown import markdown
                
                html_content = markdown(content)
                html_doc = f"""
                <!DOCTYPE html>
                <html>
                <head>
                    <meta charset="utf-8">
                    <style>
                        body {{ font-family: Arial, sans-serif; margin: 40px; line-height: 1.6; }}
                        h1 {{ color: #333; border-bottom: 2px solid #4472C4; }}
                        h2 {{ color: #555; }}
                        code {{ background: #f4f4f4; padding: 2px 6px; border-radius: 3px; }}
                        pre {{ background: #f4f4f4; padding: 10px; border-radius: 5px; }}
                    </style>
                </head>
                <body>
                    {html_content}
                </body>
                </html>
                """
                
                HTML(string=html_doc).write_pdf(str(filepath))
            
            elif format == "docx":
                # Convert Markdown to Word using python-docx
                from docx import Document
                from docx.shared import Pt, RGBColor
                
                doc = Document()
                
                # Parse Markdown and add to document
                lines = content.split('\n')
                for line in lines:
                    if line.startswith('# '):
                        doc.add_heading(line[2:], level=1)
                    elif line.startswith('## '):
                        doc.add_heading(line[3:], level=2)
                    elif line.startswith('### '):
                        doc.add_heading(line[4:], level=3)
                    elif line.strip():
                        doc.add_paragraph(line)
                
                doc.save(str(filepath))
            
            # Get file size
            size_bytes = filepath.stat().st_size
            
            # Upload to S3
            download_url = await self._upload_to_s3(filepath, user_id, session_id)
            
            generated_file = GeneratedFile(
                filename=filename,
                filepath=str(filepath),
                file_type=format,
                size_bytes=size_bytes,
                download_url=download_url,
                created_at=datetime.now(),
            )
            
            logger.info(f"Generated document: {filename} ({size_bytes} bytes)")
            
            return FileGenerationResult(
                success=True,
                files=[generated_file],
            )
            
        except Exception as e:
            logger.error(f"Document generation failed: {e}")
            return FileGenerationResult(
                success=False,
                files=[],
                error=str(e),
            )
    
    async def generate_code_project(
        self,
        user_id: str,
        session_id: str,
        files: Dict[str, str],  # filename -> content
        project_name: Optional[str] = None,
    ) -> FileGenerationResult:
        """
        Generate a multi-file code project.
        
        Args:
            user_id: User ID
            session_id: Session ID
            files: Dict of filename -> content
            project_name: Project folder name
        
        Returns:
            FileGenerationResult with folder and zip
        """
        try:
            # Create session workspace
            session_dir = self.workspace_root / user_id / session_id
            session_dir.mkdir(parents=True, exist_ok=True)
            
            # Generate project name
            if not project_name:
                project_name = f"project_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
            
            project_dir = session_dir / project_name
            project_dir.mkdir(parents=True, exist_ok=True)
            
            # Write all files
            generated_files = []
            for filename, content in files.items():
                # Create subdirectories if needed
                file_path = project_dir / filename
                file_path.parent.mkdir(parents=True, exist_ok=True)
                
                # Write file
                file_path.write_text(content, encoding='utf-8')
                
                generated_files.append(GeneratedFile(
                    filename=filename,
                    filepath=str(file_path),
                    file_type="code",
                    size_bytes=file_path.stat().st_size,
                    download_url="",  # Individual files don't get URLs
                    created_at=datetime.now(),
                ))
            
            # Create zip file
            zip_filename = f"{project_name}.zip"
            zip_path = session_dir / zip_filename
            
            with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                for file in project_dir.rglob('*'):
                    if file.is_file():
                        zipf.write(file, file.relative_to(project_dir))
            
            # Upload zip to S3
            download_url = await self._upload_to_s3(zip_path, user_id, session_id)
            
            logger.info(f"Generated code project: {project_name} ({len(files)} files)")
            
            return FileGenerationResult(
                success=True,
                files=generated_files,
                folder_path=str(project_dir),
                zip_path=str(zip_path),
            )
            
        except Exception as e:
            logger.error(f"Code project generation failed: {e}")
            return FileGenerationResult(
                success=False,
                files=[],
                error=str(e),
            )
    
    async def _upload_to_s3(self, filepath: Path, user_id: str, session_id: str) -> str:
        """
        Upload file to S3 and return public URL.
        
        Args:
            filepath: Local file path
            user_id: User ID
            session_id: Session ID
        
        Returns:
            Public download URL
        """
        if not self.s3_client:
            # Return local path if S3 not available
            return f"/api/files/{user_id}/{session_id}/{filepath.name}"
        
        try:
            # TODO: Implement S3 upload
            # s3_key = f"generated/{user_id}/{session_id}/{filepath.name}"
            # self.s3_client.upload_file(str(filepath), BUCKET_NAME, s3_key)
            # return f"https://{BUCKET_NAME}.s3.amazonaws.com/{s3_key}"
            
            return f"/api/files/{user_id}/{session_id}/{filepath.name}"
        except Exception as e:
            logger.error(f"S3 upload failed: {e}")
            return f"/api/files/{user_id}/{session_id}/{filepath.name}"
    
    async def cleanup_session(self, user_id: str, session_id: str):
        """Clean up session workspace."""
        session_dir = self.workspace_root / user_id / session_id
        if session_dir.exists():
            shutil.rmtree(session_dir)
            logger.info(f"Cleaned up session workspace: {session_id}")


# Global instance
_file_generation_service: Optional[FileGenerationService] = None


def get_file_generation_service() -> FileGenerationService:
    """Get or create the global file generation service."""
    global _file_generation_service
    if _file_generation_service is None:
        _file_generation_service = FileGenerationService()
    return _file_generation_service
