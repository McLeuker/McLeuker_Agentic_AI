"""
Document Analyzer - Document Processing and Analysis
=====================================================

Analyzes various document types:
- PDFs
- Word documents
- Spreadsheets
- Presentations
- Text files
"""

import json
import logging
from typing import Dict, List, Optional, Any
from pathlib import Path
import mimetypes

logger = logging.getLogger(__name__)


class DocumentAnalyzer:
    """
    Analyzes documents using kimi-2.5.
    
    Supports:
    - PDF analysis
    - Word document analysis
    - Spreadsheet analysis
    - Text extraction
    - Summarization
    """
    
    def __init__(self, llm_client):
        self.llm_client = llm_client
    
    async def analyze_document(
        self,
        file_path: str,
        query: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Analyze a document.
        
        Args:
            file_path: Path to document
            query: Optional specific query
            
        Returns:
            Analysis result
        """
        path = Path(file_path)
        mime_type, _ = mimetypes.guess_type(str(path))
        
        if mime_type == "application/pdf":
            return await self._analyze_pdf(file_path, query)
        elif mime_type in [
            "application/msword",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ]:
            return await self._analyze_word(file_path, query)
        elif mime_type in [
            "application/vnd.ms-excel",
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        ]:
            return await self._analyze_spreadsheet(file_path, query)
        elif mime_type in ["text/plain", "text/markdown"]:
            return await self._analyze_text(file_path, query)
        else:
            return {"error": f"Unsupported document type: {mime_type}"}
    
    async def _analyze_pdf(
        self,
        file_path: str,
        query: Optional[str]
    ) -> Dict[str, Any]:
        """Analyze PDF document"""
        
        try:
            # Try to import PyPDF2
            try:
                import PyPDF2
            except ImportError:
                return {
                    "error": "PyPDF2 not installed. Install with: pip install PyPDF2",
                    "suggestion": "For now, you can convert PDF to text manually"
                }
            
            # Extract text
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n\n"
            
            if not text.strip():
                return {"error": "Could not extract text from PDF"}
            
            # Analyze with kimi-2.5
            user_prompt = query or "Analyze this document. Provide a summary, key points, and main topics."
            
            messages = [
                {
                    "role": "system",
                    "content": "You are a document analysis expert. Extract key information and provide structured analysis."
                },
                {
                    "role": "user",
                    "content": f"{user_prompt}\n\nDocument content:\n{text[:10000]}"
                }
            ]
            
            response = await self.llm_client.chat.completions.create(
                model="kimi-k2.5",
                messages=messages,
                temperature=1.0,
                max_tokens=4000
            )
            
            analysis = response.choices[0].message.content
            
            return {
                "document_type": "PDF",
                "pages": len(reader.pages),
                "analysis": analysis,
                "text_preview": text[:1000] + "..." if len(text) > 1000 else text
            }
            
        except Exception as e:
            logger.error(f"PDF analysis failed: {e}")
            return {"error": str(e)}
    
    async def _analyze_word(
        self,
        file_path: str,
        query: Optional[str]
    ) -> Dict[str, Any]:
        """Analyze Word document"""
        
        try:
            try:
                import docx
            except ImportError:
                return {
                    "error": "python-docx not installed. Install with: pip install python-docx"
                }
            
            doc = docx.Document(file_path)
            
            # Extract text
            text = []
            for para in doc.paragraphs:
                text.append(para.text)
            
            full_text = "\n".join(text)
            
            # Analyze
            user_prompt = query or "Analyze this document. Provide a summary and key points."
            
            messages = [
                {
                    "role": "system",
                    "content": "You are a document analysis expert."
                },
                {
                    "role": "user",
                    "content": f"{user_prompt}\n\nDocument content:\n{full_text[:10000]}"
                }
            ]
            
            response = await self.llm_client.chat.completions.create(
                model="kimi-k2.5",
                messages=messages,
                temperature=1.0,
                max_tokens=4000
            )
            
            return {
                "document_type": "Word",
                "paragraphs": len(doc.paragraphs),
                "analysis": response.choices[0].message.content
            }
            
        except Exception as e:
            logger.error(f"Word analysis failed: {e}")
            return {"error": str(e)}
    
    async def _analyze_spreadsheet(
        self,
        file_path: str,
        query: Optional[str]
    ) -> Dict[str, Any]:
        """Analyze spreadsheet"""
        
        try:
            try:
                import openpyxl
            except ImportError:
                return {
                    "error": "openpyxl not installed. Install with: pip install openpyxl"
                }
            
            wb = openpyxl.load_workbook(file_path, data_only=True)
            
            # Get summary of sheets
            sheets_info = []
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                sheets_info.append({
                    "name": sheet_name,
                    "rows": sheet.max_row,
                    "columns": sheet.max_column
                })
            
            # Extract sample data from first sheet
            first_sheet = wb[wb.sheetnames[0]]
            data_sample = []
            for row in first_sheet.iter_rows(min_row=1, max_row=min(10, first_sheet.max_row), values_only=True):
                data_sample.append(row)
            
            return {
                "document_type": "Spreadsheet",
                "sheets": sheets_info,
                "sample_data": data_sample,
                "note": "For detailed analysis, convert to CSV or provide specific questions"
            }
            
        except Exception as e:
            logger.error(f"Spreadsheet analysis failed: {e}")
            return {"error": str(e)}
    
    async def _analyze_text(
        self,
        file_path: str,
        query: Optional[str]
    ) -> Dict[str, Any]:
        """Analyze text file"""
        
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
            
            user_prompt = query or "Analyze this document. Provide a summary and key points."
            
            messages = [
                {
                    "role": "system",
                    "content": "You are a document analysis expert."
                },
                {
                    "role": "user",
                    "content": f"{user_prompt}\n\nDocument content:\n{text[:10000]}"
                }
            ]
            
            response = await self.llm_client.chat.completions.create(
                model="kimi-k2.5",
                messages=messages,
                temperature=1.0,
                max_tokens=4000
            )
            
            return {
                "document_type": "Text",
                "character_count": len(text),
                "analysis": response.choices[0].message.content
            }
            
        except Exception as e:
            logger.error(f"Text analysis failed: {e}")
            return {"error": str(e)}
    
    async def summarize_document(
        self,
        file_path: str,
        max_length: str = "medium"
    ) -> Dict[str, Any]:
        """
        Summarize a document.
        
        Args:
            file_path: Path to document
            max_length: Summary length (short, medium, long)
            
        Returns:
            Summary result
        """
        lengths = {
            "short": "1-2 paragraphs",
            "medium": "3-5 paragraphs",
            "long": "comprehensive summary"
        }
        
        query = f"Provide a {lengths.get(max_length, 'medium')} summary of this document."
        return await self.analyze_document(file_path, query)
    
    async def extract_key_information(
        self,
        file_path: str,
        information_types: List[str]
    ) -> Dict[str, Any]:
        """
        Extract specific types of information from document.
        
        Args:
            file_path: Path to document
            information_types: Types of information to extract
            
        Returns:
            Extracted information
        """
        query = f"Extract the following information: {', '.join(information_types)}"
        return await self.analyze_document(file_path, query)
