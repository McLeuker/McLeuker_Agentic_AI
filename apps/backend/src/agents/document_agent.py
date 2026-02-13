"""
Document Agent — Professional Document Generation (DOCX, PDF)
==============================================================

Creates professional documents from natural language descriptions.
Supports templates, formatting, and rich content.

Features:
- Template-based document generation
- Rich formatting (styles, tables, images)
- Multiple output formats (DOCX, PDF)
- Content structuring and organization
"""

import json
import logging
import os
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, AsyncGenerator, Dict, List, Optional
import openai

# Document generation libraries
try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not available - DOCX generation disabled")

try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logging.warning("reportlab not available - PDF generation disabled")

logger = logging.getLogger(__name__)


@dataclass
class DocumentSection:
    """A section in a document."""
    title: str
    content: str
    level: int = 1  # Heading level (1-6)
    style: str = "normal"
    bullet_points: List[str] = field(default_factory=list)
    table_data: Optional[List[List[str]]] = None
    
    def to_dict(self) -> Dict:
        return {
            "title": self.title,
            "content": self.content,
            "level": self.level,
            "style": self.style,
            "bullet_points": self.bullet_points,
            "table_data": self.table_data,
        }


@dataclass
class DocumentTemplate:
    """Template for document generation."""
    name: str
    description: str
    sections: List[str]
    default_styles: Dict[str, Any]
    
    def to_dict(self) -> Dict:
        return {
            "name": self.name,
            "description": self.description,
            "sections": self.sections,
            "default_styles": self.default_styles,
        }


@dataclass
class GeneratedDocument:
    """A generated document."""
    document_id: str
    title: str
    file_path: str
    file_type: str
    sections: List[DocumentSection]
    metadata: Dict[str, Any] = field(default_factory=dict)
    
    def to_dict(self) -> Dict:
        return {
            "document_id": self.document_id,
            "title": self.title,
            "file_path": self.file_path,
            "file_type": self.file_type,
            "sections": [s.to_dict() for s in self.sections],
            "metadata": self.metadata,
        }


class DocumentAgent:
    """
    AI Document Agent for professional document generation.
    
    Usage:
        agent = DocumentAgent(llm_client)
        async for event in agent.generate_document(
            "Create a business proposal for a software consulting company",
            file_type="docx"
        ):
            print(event)
    """
    
    # Built-in templates
    TEMPLATES = {
        "business_proposal": DocumentTemplate(
            name="Business Proposal",
            description="Professional business proposal with executive summary, scope, timeline, and pricing",
            sections=[
                "Executive Summary",
                "Company Overview",
                "Problem Statement",
                "Proposed Solution",
                "Scope of Work",
                "Timeline",
                "Pricing",
                "Terms and Conditions",
                "Contact Information",
            ],
            default_styles={
                "font": "Calibri",
                "heading_color": "#2E5090",
                "font_size": 11,
            },
        ),
        "research_report": DocumentTemplate(
            name="Research Report",
            description="Academic-style research report with abstract, methodology, findings, and conclusions",
            sections=[
                "Abstract",
                "Introduction",
                "Literature Review",
                "Methodology",
                "Findings",
                "Analysis",
                "Conclusions",
                "References",
            ],
            default_styles={
                "font": "Times New Roman",
                "heading_color": "#000000",
                "font_size": 12,
            },
        ),
        "project_plan": DocumentTemplate(
            name="Project Plan",
            description="Comprehensive project plan with objectives, tasks, resources, and milestones",
            sections=[
                "Project Overview",
                "Objectives",
                "Scope",
                "Deliverables",
                "Timeline",
                "Resources",
                "Risk Assessment",
                "Budget",
                "Communication Plan",
            ],
            default_styles={
                "font": "Arial",
                "heading_color": "#1F4E79",
                "font_size": 11,
            },
        ),
        "meeting_minutes": DocumentTemplate(
            name="Meeting Minutes",
            description="Structured meeting minutes with attendees, agenda, discussions, and action items",
            sections=[
                "Meeting Information",
                "Attendees",
                "Agenda",
                "Discussion Points",
                "Decisions",
                "Action Items",
                "Next Meeting",
            ],
            default_styles={
                "font": "Arial",
                "heading_color": "#404040",
                "font_size": 11,
            },
        ),
        "resume": DocumentTemplate(
            name="Resume/CV",
            description="Professional resume with summary, experience, education, and skills",
            sections=[
                "Contact Information",
                "Professional Summary",
                "Work Experience",
                "Education",
                "Skills",
                "Certifications",
                "References",
            ],
            default_styles={
                "font": "Calibri",
                "heading_color": "#2B579A",
                "font_size": 11,
            },
        ),
    }
    
    def __init__(
        self,
        llm_client: openai.AsyncOpenAI,
        model: str = "kimi-k2.5",
        output_dir: str = "/tmp/document_agent",
    ):
        self.llm_client = llm_client
        self.model = model
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
        # Track generated documents
        self._documents: Dict[str, GeneratedDocument] = {}
    
    async def generate_document(
        self,
        description: str,
        file_type: str = "docx",
        template: Optional[str] = None,
        title: Optional[str] = None,
        context: Optional[Dict] = None,
    ) -> AsyncGenerator[Dict[str, Any], None]:
        """
        Generate a professional document from a description.
        
        Args:
            description: Natural language description of the document
            file_type: Output format (docx, pdf)
            template: Optional template name to use
            title: Optional document title
            context: Additional context for generation
            
        Yields:
            Progress events and final document info
        """
        import uuid
        document_id = str(uuid.uuid4())
        
        yield {"type": "start", "data": {"document_id": document_id, "description": description}}
        
        try:
            # Step 1: Analyze requirements and select/confirm template
            yield {"type": "phase", "data": {"phase": "analysis", "status": "started"}}
            
            selected_template = None
            if template and template in self.TEMPLATES:
                selected_template = self.TEMPLATES[template]
            else:
                selected_template = await self._select_template(description)
            
            yield {
                "type": "phase",
                "data": {
                    "phase": "analysis",
                    "status": "completed",
                    "template": selected_template.name if selected_template else "custom",
                }
            }
            
            # Step 2: Generate content for each section
            yield {"type": "phase", "data": {"phase": "content_generation", "status": "started"}}
            
            sections = await self._generate_content(
                description,
                selected_template,
                context,
            )
            
            yield {
                "type": "phase",
                "data": {
                    "phase": "content_generation",
                    "status": "completed",
                    "sections": len(sections),
                }
            }
            
            # Step 3: Create document file
            yield {"type": "phase", "data": {"phase": "file_creation", "status": "started"}}
            
            doc_title = title or await self._generate_title(description)
            
            if file_type.lower() == "docx":
                file_path = await self._create_docx(document_id, doc_title, sections, selected_template)
            elif file_type.lower() == "pdf":
                file_path = await self._create_pdf(document_id, doc_title, sections, selected_template)
            else:
                yield {"type": "error", "data": {"message": f"Unsupported file type: {file_type}"}}
                return
            
            yield {"type": "phase", "data": {"phase": "file_creation", "status": "completed", "path": file_path}}
            
            # Step 4: Create document record
            document = GeneratedDocument(
                document_id=document_id,
                title=doc_title,
                file_path=file_path,
                file_type=file_type.lower(),
                sections=sections,
                metadata={
                    "template": selected_template.name if selected_template else "custom",
                    "description": description,
                    "context": context or {},
                },
            )
            self._documents[document_id] = document
            
            yield {
                "type": "complete",
                "data": {
                    "document_id": document_id,
                    "title": doc_title,
                    "file_path": file_path,
                    "file_type": file_type,
                }
            }
            
        except Exception as e:
            logger.error(f"Error generating document: {e}")
            yield {"type": "error", "data": {"message": str(e)}}
    
    async def _select_template(self, description: str) -> Optional[DocumentTemplate]:
        """Select the best template for the description."""
        messages = [
            {
                "role": "system",
                "content": f"""Select the best document template for this request.

Available templates:
{json.dumps({k: v.description for k, v in self.TEMPLATES.items()}, indent=2)}

Respond with JSON:
{{
    "template": "template_name or null for custom",
    "reasoning": "explanation"
}}"""
            },
            {
                "role": "user",
                "content": f"Select template for: {description}"
            }
        ]
        
        response = await self.llm_client.chat.completions.create(
            model=self.model,
            messages=messages,
            temperature=0.3,
            max_tokens=512,
            response_format={"type": "json_object"},
        )
        
        result = json.loads(response.choices[0].message.content)
        template_name = result.get("template")
        
        if template_name and template_name in self.TEMPLATES:
            return self.TEMPLATES[template_name]
        
        return None
    
    async def _generate_title(self, description: str) -> str:
        """Generate a title for the document."""
        messages = [
            {
                "role": "system",
                "content": "Generate a concise, professional title for this document. Respond with just the title."
            },
            {
                "role": "user",
                "content": description
            }
        ]
        
        response = await self.llm_client.chat.completions.create(
            model=self.model,
            messages=messages,
            temperature=0.7,
            max_tokens=100,
        )
        
        return response.choices[0].message.content.strip().strip('"')
    
    async def _generate_content(
        self,
        description: str,
        template: Optional[DocumentTemplate],
        context: Optional[Dict],
    ) -> List[DocumentSection]:
        """Generate content for all sections."""
        sections = []
        
        section_names = template.sections if template else ["Main Content"]
        
        for section_name in section_names:
            section = await self._generate_section(
                section_name,
                description,
                template,
                context,
            )
            sections.append(section)
        
        return sections
    
    async def _generate_section(
        self,
        section_name: str,
        description: str,
        template: Optional[DocumentTemplate],
        context: Optional[Dict],
    ) -> DocumentSection:
        """Generate content for a single section."""
        messages = [
            {
                "role": "system",
                "content": f"""Generate content for the "{section_name}" section of a document.

Document description: {description}

Respond with JSON:
{{
    "title": "Section title",
    "content": "Main paragraph content (2-4 paragraphs)",
    "bullet_points": ["point 1", "point 2", "point 3"],
    "table_data": null or [["col1", "col2"], ["val1", "val2"]]
}}"""
            }
        ]
        
        if context:
            messages.append({
                "role": "user",
                "content": f"Additional context: {json.dumps(context)}"
            })
        
        response = await self.llm_client.chat.completions.create(
            model=self.model,
            messages=messages,
            temperature=0.7,
            max_tokens=2048,
            response_format={"type": "json_object"},
        )
        
        result = json.loads(response.choices[0].message.content)
        
        return DocumentSection(
            title=result.get("title", section_name),
            content=result.get("content", ""),
            level=1,
            bullet_points=result.get("bullet_points", []),
            table_data=result.get("table_data"),
        )
    
    async def _create_docx(
        self,
        document_id: str,
        title: str,
        sections: List[DocumentSection],
        template: Optional[DocumentTemplate],
    ) -> str:
        """Create a DOCX file."""
        if not DOCX_AVAILABLE:
            raise RuntimeError("python-docx not available")
        
        file_path = self.output_dir / f"{document_id}.docx"
        
        doc = Document()
        
        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = template.default_styles.get("font", "Calibri") if template else "Calibri"
        font.size = Pt(template.default_styles.get("font_size", 11) if template else 11)
        
        # Add title
        title_para = doc.add_heading(title, 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add sections
        for section in sections:
            # Add section heading
            heading = doc.add_heading(section.title, level=section.level)
            
            # Add content
            if section.content:
                doc.add_paragraph(section.content)
            
            # Add bullet points
            for point in section.bullet_points:
                p = doc.add_paragraph(point, style='List Bullet')
            
            # Add table if present
            if section.table_data:
                table = doc.add_table(rows=len(section.table_data), cols=len(section.table_data[0]))
                for i, row_data in enumerate(section.table_data):
                    for j, cell_data in enumerate(row_data):
                        table.cell(i, j).text = str(cell_data)
            
            # Add spacing between sections
            doc.add_paragraph()
        
        doc.save(str(file_path))
        return str(file_path)
    
    async def _create_pdf(
        self,
        document_id: str,
        title: str,
        sections: List[DocumentSection],
        template: Optional[DocumentTemplate],
    ) -> str:
        """Create a PDF file."""
        if not PDF_AVAILABLE:
            raise RuntimeError("reportlab not available")
        
        file_path = self.output_dir / f"{document_id}.pdf"
        
        doc = SimpleDocTemplate(
            str(file_path),
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18,
        )
        
        styles = getSampleStyleSheet()
        story = []
        
        # Title style
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            spaceAfter=30,
            alignment=1,  # Center
        )
        
        # Add title
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 0.2 * inch))
        
        # Add sections
        for section in sections:
            # Section heading
            story.append(Paragraph(section.title, styles['Heading2']))
            story.append(Spacer(1, 0.1 * inch))
            
            # Content
            if section.content:
                story.append(Paragraph(section.content, styles['Normal']))
                story.append(Spacer(1, 0.1 * inch))
            
            # Bullet points
            for point in section.bullet_points:
                story.append(Paragraph(f"• {point}", styles['Normal']))
            
            if section.bullet_points:
                story.append(Spacer(1, 0.1 * inch))
            
            # Table
            if section.table_data:
                table = Table(section.table_data)
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 14),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ]))
                story.append(table)
                story.append(Spacer(1, 0.2 * inch))
            
            story.append(Spacer(1, 0.2 * inch))
        
        doc.build(story)
        return str(file_path)
    
    def list_templates(self) -> List[Dict]:
        """List available document templates."""
        return [{"id": k, **v.to_dict()} for k, v in self.TEMPLATES.items()]
    
    def get_document(self, document_id: str) -> Optional[GeneratedDocument]:
        """Get a generated document by ID."""
        return self._documents.get(document_id)


# Singleton instance
_document_agent: Optional[DocumentAgent] = None


def get_document_agent(llm_client: openai.AsyncOpenAI = None) -> Optional[DocumentAgent]:
    """Get or create the Document Agent singleton."""
    global _document_agent
    if _document_agent is None and llm_client:
        _document_agent = DocumentAgent(llm_client)
    return _document_agent
